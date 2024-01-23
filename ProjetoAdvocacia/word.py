'''
from docx import Document

def encontrar_e_armazenar_linhas(caminho_arquivo, palavras_chave):
    doc = Document(caminho_arquivo)
    linhas_encontradas = []
    valores = []
    for tabela in doc.tables:
        for linha in tabela.rows:
            # Convertendo cada célula para minúsculas
            dados_linha = [celula.text.strip().lower() for celula in linha.cells]
            #print(dados_linha)
            # Verifica se alguma palavra-chave está presente na linha (insensível a maiúsculas e minúsculas)
            for palavra in palavras_chave:
                if palavra in dados_linha[3]:
                    linhas_encontradas.append(dados_linha)
                    tamanho = len(dados_linha)
                    #print(palavra)
                    valores.append(dados_linha[tamanho-1])

    return linhas_encontradas, valores
'''
from docx import Document

def dados_word(arquivo_selecionado, palavras_chave):
    doc = Document('static/arquivos/'+arquivo_selecionado)
    linhas_encontradas = []
    valores = []
    for tabela in doc.tables:
        for linha in tabela.rows:
            # Convertendo cada célula para minúsculas
            dados_linha = [celula.text.strip().lower() for celula in linha.cells]
            #print(dados_linha)
            # Verifica se alguma palavra-chave está presente na linha (insensível a maiúsculas e minúsculas)
            for palavra in palavras_chave:
                if palavra in dados_linha[3]:
                    linhas_encontradas.append(dados_linha)
                    tamanho = len(dados_linha)
                    #print(palavra)
                    valores.append(dados_linha[tamanho-1])
                    dicionario = {  'palavra_chave':'exemplo',
                            'total':'total',
                            #'informacoes': 'dicionario={'pagina':pagina,'valor':valor}'
                            }

    return linhas_encontradas, valores